from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.text.run import Font, Run
from docx.section import Section
import zipfile
import xml.etree.cElementTree as ET
import xml.dom.minidom


document = Document()

txt1 = "TEXT1"
txt2 = "Text2"
txt3 = "Text3"
txt4 = "Text4"
txt5 = "Text5"
txt6 = "Text6"
txt7 = "Text7"
txt8 = "Text8"
txt9 = "Text9"
img1 = "img1.png"
img2 = "img2.png"
strabc = "abc"
strdef = "def"

img1Align = document.add_picture(img1)
last_paragraph = document.paragraphs[-1] 
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


document.add_picture(img2)
head1 = document.add_heading(txt1, level=0)

head1.alignment = WD_ALIGN_PARAGRAPH.CENTER

formatHead1 = head1.paragraph_format
formatHead1.space_after = Pt(40)

head2 = document.add_heading(txt2+":", level=4)
head2.alignment = WD_ALIGN_PARAGRAPH.CENTER

txt3Align = document.add_paragraph(txt3)
txt3Align.alignment = WD_ALIGN_PARAGRAPH.CENTER
txt3Align.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE



head3 = document.add_heading(txt4+":", level=4)
head3.alignment = WD_ALIGN_PARAGRAPH.CENTER

txt5Align = document.add_paragraph(txt5)
txt5Align.add_run().font.highlight_color = WD_COLOR_INDEX.YELLOW

txt5Align.alignment = WD_ALIGN_PARAGRAPH.CENTER
txt5Align.line_spacing_rule = WD_LINE_SPACING.DOUBLE

head4 = document.add_heading(txt6+":", level=4)
head4.alignment = WD_ALIGN_PARAGRAPH.CENTER

txt7Align = document.add_paragraph(txt7)
txt7Align.alignment = WD_ALIGN_PARAGRAPH.CENTER
formatTxtAlign7 = txt7Align.paragraph_format
formatTxtAlign7.space_after = Pt(40)


head5 = document.add_heading(txt8+": "+strabc, level=4)
head5.alignment = WD_ALIGN_PARAGRAPH.CENTER


head5 = document.add_heading(txt9+": "+strdef, level=4)
head5.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_page_break()


dc = document.save("demo.docx")



